import pyautogui
from docx import Document
from docx.shared import Cm, Pt
from datetime  import date
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
import time 
from qt3 import Ui_MainWindow
import os
#from tkinter import Tk
#from tkinter.filedialog import askdirectory, asksaveasfilename

    

xa = 1
xb = 1

# xb = int(random()*10000)
today = date.today()



class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self, *args, **kwargs):
        super(MainWindow, self).__init__(*args, **kwargs)
       
        

        self.setupUi(self)
        self.btn_menu_1.pressed.connect(self.starter)
        self.btn_menu_4.pressed.connect(self.quit)
        self.btn_menu_3.pressed.connect(self.done)
        self.btn_menu_2.pressed.connect(self.myscreentaker)
        self.lineEdit.returnPressed.connect(self.btn_add_note.click)
        self.btn_add_note.pressed.connect(self.take_notes)
        self.btn_info.pressed.connect(self.info)
        self.btn_wth.pressed.connect(self.about)
        self.counter = xa
        self.xls_name = xb  
        self.issaved = "y"      
        self.show()

    def starter (self):
        try:
            self.doc = Document()
            self.run = self.doc.add_paragraph().add_run('Kontrol Tarihi    : {} \nKontrol Sonucu : Başarılı / Başarısız'.format(today.strftime("%d.%m.%Y")))
            self.font = self.run.font
            self.font.name = 'Calibri'
            self.font.size = Pt(12) 
            if  not os.path.exists("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name)):
                os.makedirs("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name))
            else:
                while  os.path.exists("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name)):
                    self.xls_name = self.xls_name + 1        
                    if  not os.path.exists("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name)):
                        os.makedirs("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name))
                        break
            
            self.label.setText("<html><head/><body><p><span style=\" color:#00CED1;\">{}_WTH_{} file is created.</span></p></body></html>".format(today.strftime("%d%m%Y"),self.xls_name))            
            self.issaved = "n" 
        except:
            self.label.setText("<html><head/><body><p><span style=\" color: rgb(255, 30, 30);\">Error: The request could not be satisfied.</span></p></body></html>")    
        


    def myscreentaker(self):
        try:   
            myScreenshot = pyautogui.screenshot()
            myScreenshot.save('{}/file_{}.png'.format("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name), self.counter))
            self.doc.add_picture('{}/file_{}.png'.format("{}_WTH_{}".format(today.strftime("%d%m%Y"),self.xls_name), self.counter), width=Cm(15.24), height=Cm(8.57))
            self.label.setText("<html><head/><body><p><span style=\" color:#00CED1;\">Number of ScreenShots: {}</span></p></body></html>".format(self.counter))  
            self.counter = self.counter +1
            self.issaved = "n" 
        except:
            self.label.setText("<html><head/><body><p><span style=\" color: rgb(255, 30, 30);\">Error:Make sure you hit the Start button.</span></p></body></html>")    
        

    def take_notes(self):
        try:
            self.run = self.doc.add_paragraph().add_run("\n" + self.lineEdit.text())
            self.font = self.run.font
            self.font.name = 'Calibri'
            self.font.size = Pt(12)
            self.label.setText("<html><head/><body><p><span style=\" color:#00CED1;\">Your note is added successfully.</span></p></body></html>")    
            self.lineEdit.clear()
            self.issaved = "n" 
        except:
            self.label.setText("<html><head/><body><p><span style=\" color: rgb(255, 30, 30);\">Error:Your note could not be added.</span></p></body></html>")    
        



    def done(self):
        #root=Tk()
        #root.withdraw()
        #fout = asksaveasfilename(filetypes=[('Word','.docx')], defaultextension = ".docx")
        
        fout = QFileDialog.getSaveFileName(self,'Save file', '', "Word File (*.docx)")

        try:
            self.doc.save("{}".format(fout[0]))
            #self.doc.save("{}_WTH_{}.docx".format(today.strftime("%d%m%Y"),self.xls_name))
            self.label.setText("<html><head/><body><p><span style=\" color:#00CED1;\">Successfully saved | {} </span></p></body></html>" .format(fout[0]))
            self.issaved = "y" 
        except:
            self.label.setText("<html><head/><body><p><span style=\" color: rgb(255, 30, 30);\">Error:Your changes could not be saved.</span></p></body></html>")    
        

    def info(self):
        self.label.setText(("<html><head/><body><p><span style=\" color: rgb(119, 221, 0);\">CTRL+S for Screen Shot | Press Enter for taking notes</span></p></body></html>"))
 

    def about(self):
        self.label.setText(("<html><head/><body><p><span style=\" color: rgb(119, 221, 0);\">Developed by okinvan on behalf of IRM Team.</span></p></body></html>"))


    def quit(self):
        
        if self.issaved=="n":
            choice = QMessageBox.question(self, 'Warning!',
            "<html><head/><body><p><span style=\" color: rgb(119, 221, 0);\">There are unsaved changes. Do you want to save before leaving?</span></p></body></html>"
            
            , QMessageBox.Yes | QMessageBox.No)
            if choice == QMessageBox.Yes: 
                
                self.done()
            else:
                app.quit()    
          
        else:
            app.quit()
        




if __name__ == '__main__':
    app = QApplication([])

    app_icon = QIcon()
    app_icon.addFile('x3.ico', QSize(48,48))
    app.setWindowIcon(app_icon)

    app.setStyle('Fusion')
    app.setApplicationName("WTH")
    
    window = MainWindow()
    app.exec_()

  